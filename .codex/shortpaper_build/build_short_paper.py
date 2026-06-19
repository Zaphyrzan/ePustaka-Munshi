from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(
    r"C:\Users\wanza\OneDrive - Universiti Teknologi Malaysia (UTM)\UTM\Final Year Project\FINAL"
    r"\ePustaka Munshi Short Paper - Draft.docx"
)

FONT = "Times New Roman"
BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(80, 80, 80)
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
PLACEHOLDER_FILL = "F7F7F7"
CONTENT_WIDTH_IN = 6.5


def set_run_font(run, size=10, bold=None, italic=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def set_columns(section, num=1, space=360):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space))


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_styles(doc):
    return
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, before, after, color in [
        ("Heading 1", 10.5, 8, 4, BLUE),
        ("Heading 2", 10, 5, 3, BLUE),
        ("Heading 3", 9.5, 4, 2, BLUE),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Caption" in styles:
        cap = styles["Caption"]
    else:
        cap = styles.add_style("Caption", 1)
    cap.font.name = FONT
    cap._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    cap.font.size = Pt(8.5)
    cap.font.italic = True
    cap.font.color.rgb = MUTED
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(5)

    ref = styles.add_style("Reference Entry", 1)
    ref.font.name = FONT
    ref._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    ref._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    ref.font.size = Pt(8.5)
    ref.paragraph_format.left_indent = Inches(0.25)
    ref.paragraph_format.first_line_indent = Inches(-0.25)
    ref.paragraph_format.space_after = Pt(3)
    ref.paragraph_format.line_spacing = 1.0


def add_para(doc, text="", size=10, bold=False, italic=False, align=None, color=None, after=None, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, size=10.5 if level == 1 else 10, bold=True, color=BLUE)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    r = p.add_run(text)
    set_run_font(r, size=8.5, italic=True, color=MUTED)
    return p


def add_placeholder(doc, label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    short_label = label.split(".")[0] if label.startswith("Figure ") else label
    r = p.add_run(f"[IMAGE HERE - {short_label}]")
    set_run_font(r, size=10, bold=True, color=MUTED)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    add_caption(doc, label)


def add_table(doc, caption, headers, rows, widths):
    add_caption(doc, caption)
    header_line = " | ".join(headers)
    p = add_para(doc, header_line, size=8.8, bold=True, after=2)
    p.paragraph_format.left_indent = Inches(0.18)
    for row in rows:
        if len(row) == 2:
            line = f"{row[0]}: {row[1]}"
        else:
            line = " | ".join(row)
        p = add_para(doc, line, size=8.8, after=2)
        p.paragraph_format.left_indent = Inches(0.18)
    add_para(doc, "", after=2)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        set_run_font(run, size=9.5)
    r = p.add_run(text)
    set_run_font(r, size=9.5)
    return p


doc = Document()
configure_styles(doc)
for section in doc.sections:
    configure_section(section)
    set_columns(section, 1)

section = doc.sections[0]
add_para(doc, "International Journal of Innovative Computing XX(nn) pp-pp", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, after=20)
add_para(doc, "EPUSTAKA MUNSHI: SMART LIBRARY SYSTEM WITH OCR LEDGER DIGITIZATION", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
add_para(doc, "WAN ZAFIRZAN BIN WAN TARMIZAN", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
add_para(doc, "Malaysia-Japan International Institute of Technology, Universiti Teknologi Malaysia", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
add_para(doc, "Bachelor of Computer Science (Software Engineering) | A22MJ8003", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
add_para(doc, "Supervisor: Dr. Siti Nur Khadijah Aishah Binti Ibrahim", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

abstract = (
    "Abstract - Secondary-school libraries that still depend on handwritten ledgers face persistent problems in catalog search, "
    "book-status tracking, borrowing control, and preservation of historical acquisition records. This paper presents ePustaka "
    "Munshi, a web-based smart library system developed for SMK Abdullah Munshi to modernize a paper-based library workflow. "
    "The system integrates book and copy cataloguing, student and member management, role-based staff access, barcode-based "
    "checkout and return, bilingual user interfaces, and an OCR-assisted workflow for digitizing handwritten Buku Induk records. "
    "A React frontend, Flask REST API, and Supabase PostgreSQL database form the main three-tier architecture, while Anthropic "
    "Claude Vision is used as the primary handwritten-ledger extraction engine with Tesseract retained as a traditional OCR "
    "baseline and page-orientation aid. The system was developed using Agile Scrum across four functional sprints: user and role "
    "management, catalog and inventory, ledger digitization, and circulation. Testing covered functional validation, integration "
    "checks, OCR review correction, and user acceptance feedback from the school library context. The completed system reduces "
    "dependence on fragile paper records, provides a searchable digital catalogue, supports faster barcode circulation, and "
    "preserves historical ledger data through a human-verified OCR process."
)
p = add_para(doc, abstract, size=9.5, bold=False, after=5)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw = add_para(
    doc,
    "Keywords - Smart library system; OCR ledger digitization; handwritten document transcription; barcode circulation; React; Flask; Supabase.",
    size=9.2,
    italic=True,
    after=8,
)
kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_heading(doc, "I. INTRODUCTION", 1)
for text in [
    (
        "Library resource management is a practical requirement for supporting reading culture, academic reference use, and "
        "day-to-day school administration. At SMK Abdullah Munshi in Penang, Malaysia, library operations had long depended on "
        "physical ledgers, manual borrowing logs, and simple spreadsheet records. The school library manages more than 7,000 "
        "books, yet important fields such as titles, authors, publishers, Dewey Decimal Classification, acquisition numbers, "
        "and loan statuses were fragmented across physical and digital records."
    ),
    (
        "The manual process created three connected problems. First, the physical ledger was difficult to search and vulnerable "
        "to deterioration. Second, borrowing and returning required repeated manual writing, especially during peak school hours. "
        "Third, previous digitization attempts did not provide validation, duplication checking, real-time availability, or "
        "role-based permissions. Similar school-library modernization challenges have been reported in web-based library "
        "system studies, where paper ledgers restrict searchability and increase record-loss risk [1], [6]."
    ),
    (
        "This project therefore developed ePustaka Munshi, a web-based library management system with OCR ledger digitization. "
        "The goal is not merely to replace a paper ledger with a database, but to create a workflow that fits the real operating "
        "conditions of a Malaysian secondary-school library: teachers supervise, library prefects help at the counter, students "
        "need simple access to catalogue and loan information, and old handwritten records must be preserved before they degrade."
    ),
]:
    p = add_para(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_heading(doc, "II. BACKGROUND AND RELATED WORK", 1)
for text in [
    (
        "Digital transformation in school libraries is increasingly linked to collection visibility, administrative efficiency, "
        "and equitable access to learning resources. Studies on digital library management and digital collection strategies "
        "emphasize that school libraries benefit when metadata, circulation, and user records are centralized rather than kept "
        "in isolated physical files [1], [11]."
    ),
    (
        "Existing library platforms were evaluated against the needs of SMK Abdullah Munshi. KOHA provides mature integrated "
        "library functionality, but its server administration and cataloguing complexity make it heavy for a small secondary-school "
        "deployment. Legacy desktop systems such as e-Kutub Khanah match older Malaysian school contexts, but are vulnerable to "
        "device dependency, compatibility issues, and data-loss risk. Spreadsheet-based approaches are simple, but they do not "
        "provide circulation automation, permission control, or OCR-assisted migration of handwritten records."
    ),
    (
        "OCR research shows why the project required a hybrid recognition strategy. Classical OCR remains useful for printed text "
        "and image preprocessing [9], [13], [15], but handwritten historical documents introduce variability in ink, layout, "
        "language, and writing style [8], [16]. Recent studies suggest that vision-capable large language models can transcribe "
        "and structure handwritten historical records more effectively than conventional OCR in some archival settings [4], and "
        "vision models have also been quantitatively evaluated for OCR-like tasks [12]. These findings support the decision to "
        "use Claude Vision for handwritten ledger extraction while retaining Tesseract as a baseline and page-orientation aid."
    ),
]:
    p = add_para(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_table(
    doc,
    "Table I: Comparison of Candidate Library Solutions",
    ["Solution", "Suitability for SMK Abdullah Munshi"],
    [
        ("KOHA", "Powerful open-source ILS, but too complex for a small school library without dedicated server administration."),
        ("e-Kutub Khanah", "Localized school-library tool, but legacy desktop dependence caused maintainability and compatibility concerns."),
        ("Excel or Google Sheets", "Low barrier to entry, but lacks validation, status tracking, audit trail, and integrated circulation."),
        ("ePustaka Munshi", "Tailored web system with cloud database, OCR ledger digitization, barcode circulation, and role-based access."),
    ],
    [1.0, 2.0],
)

add_heading(doc, "III. METHODOLOGY", 1)
for text in [
    (
        "Agile Scrum was adopted because the project involved uncertain OCR performance, evolving stakeholder expectations, and "
        "workflow details that became clearer only after prototype testing with real library materials. Instead of treating OCR, "
        "circulation, and access control as a single fixed design, the project was organized into four sprints: user and role "
        "management, book catalog and inventory, ledger digitization, and borrowing or returning management."
    ),
    (
        "Requirements were collected through stakeholder discussion with the school librarian and review of existing artifacts, "
        "including the Buku Induk ledger, borrowing records, and spreadsheet inventory. The requirements were then formalized into "
        "software requirements and design documentation, including use cases, workflow diagrams, entity relationship design, test "
        "cases, and interface designs."
    ),
]:
    p = add_para(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_placeholder(doc, "Figure 1. Agile Scrum Development Lifecycle")

add_table(
    doc,
    "Table II: Technology Stack and Purpose",
    ["Component", "Purpose"],
    [
        ("React + TypeScript", "Responsive single-page user interface for catalogue, circulation, OCR review, and dashboards."),
        ("Flask REST API", "Application logic, authentication flow, permission checks, and data validation."),
        ("Supabase PostgreSQL", "Centralized relational database with managed cloud persistence and backup support [14]."),
        ("Claude Vision + Tesseract", "Handwritten ledger extraction, classical OCR comparison, and page-orientation support."),
        ("USB barcode scanner", "Keyboard-style HID input for fast checkout and return using accession or copy barcode data."),
        ("Vercel", "Web deployment target for the React frontend and Flask API entry point."),
    ],
    [1.05, 1.95],
)

add_heading(doc, "IV. SYSTEM DEVELOPMENT", 1)
add_heading(doc, "A. Requirements Analysis", 2)
for text in [
    (
        "The system requirements were derived from the current manual workflow. The library needed a secure login mechanism, "
        "centralized book and copy records, student and staff account management, role-based permissions, barcode-based loan "
        "processing, and a digitization workflow for old handwritten acquisitions. Non-functional needs included browser access, "
        "low maintenance cost, fast interaction at the circulation counter, data reliability, and usability for non-technical staff."
    ),
    (
        "The most distinctive requirement is the OCR verification workflow. Because handwritten ledger extraction is imperfect by "
        "nature, the system does not automatically merge OCR output into the catalogue. Instead, it shows extracted rows to staff, "
        "flags confidence values, allows field-by-field correction, and commits only reviewed records into the main catalogue."
    ),
]:
    p = add_para(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_placeholder(doc, "Figure 2. SMK Abdullah Munshi Handwritten Ledger")

add_heading(doc, "B. Functional and Non-Functional Requirements", 2)
add_table(
    doc,
    "Table III: Core Requirements Summary",
    ["Category", "Requirement"],
    [
        ("Authentication", "Users must log in with valid credentials and be routed to the correct role-based interface."),
        ("Inventory", "Librarians can add, edit, search, filter, and manage book and copy records."),
        ("Digitization", "Staff can upload scanned ledger pages, review OCR rows, correct fields, and commit verified records."),
        ("Circulation", "Authorized staff and prefects can scan barcodes to record borrowing, returning, renewal, and overdue status."),
        ("Performance", "Common interactions such as login and counter operations should remain fast enough for school use."),
        ("Reliability", "Unverified or low-confidence OCR data must not be inserted into the permanent catalogue without review."),
    ],
    [1.0, 2.0],
)

add_heading(doc, "C. Use Case Design", 2)
p = add_para(
    doc,
    (
        "The use case model defines four main actors: Administrator, Librarian, Library Prefect, and Student. Administrators manage "
        "global settings, staff accounts, user roles, and high-risk operations. Librarians manage book records, students, OCR jobs, "
        "and circulation monitoring. Library Prefects assist at the counter by scanning books for borrowing and returning, while "
        "students use the portal to view catalogues, active loans, overdue status, and borrowing history."
    ),
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_placeholder(doc, "Figure 3. Use Case Diagram for ePustaka Munshi")

add_heading(doc, "V. SYSTEM ARCHITECTURE DESIGN", 1)
for text in [
    (
        "The implementation follows a three-tier web architecture. The presentation layer is a React single-page application with "
        "module-based pages for authentication, catalogue browsing, circulation, user management, OCR review, and student access. "
        "The application layer is a Flask REST API that handles validation, session-aware authentication, business rules, and role "
        "permissions. The data layer is Supabase PostgreSQL, which stores users, roles, members, books, copies, loans, OCR jobs, "
        "OCR results, and digitized ledger archives."
    ),
    (
        "This architecture was selected to avoid the failure mode of a single-computer desktop system. A hosted relational database "
        "allows library data to remain accessible even if a local school computer fails. At the same time, the heaviest OCR work is "
        "kept on a controlled scanning station because handwritten recognition requires native imaging tools, API keys, and careful "
        "human review before data is merged into production records."
    ),
]:
    p = add_para(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_placeholder(doc, "Figure 4. High-Level System Workflow Swimlane Diagram")

add_heading(doc, "VI. DATABASE DESIGN", 1)
p = add_para(
    doc,
    (
        "The database design uses relational entities to separate access control, membership, bibliographic records, physical copies, "
        "loans, and OCR evidence. Roles define permission groups, while users and members represent staff and student identities. "
        "Books store bibliographic metadata, and book copies store physical inventory details such as accession number, barcode, "
        "condition, status, and shelf location. Loan records connect members, copies, handlers, due dates, renewals, and status "
        "changes. OCR jobs and OCR results preserve extraction evidence, while the digitized ledger archive keeps the original "
        "ledger fields for traceability."
    ),
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_placeholder(doc, "Figure 5. Entity Relationship Diagram for ePustaka Munshi")

add_heading(doc, "VII. INTERFACE DESIGN", 1)
for text in [
    (
        "Interface design prioritizes clarity at the school counter. The dashboard summarizes total books, active loans, overdue "
        "items, and pending OCR review. The catalogue interface supports search and filtering, while staff views provide editing "
        "controls only to users with the correct permissions."
    ),
    (
        "The OCR interface is designed as a verification workspace rather than a blind import screen. Staff upload a ledger image or "
        "PDF, review the extracted rows, correct values against the source image, and then commit selected rows. The circulation "
        "interface keeps the barcode field focused so a USB scanner behaves like direct keyboard input. Visual feedback confirms "
        "successful checkout or return and warns staff when a book is unavailable, overdue, or already processed."
    ),
]:
    p = add_para(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_placeholder(doc, "Figure 6. OCR Digitization Review Interface")
add_placeholder(doc, "Figure 7. Barcode Checkout and Return Interface")

add_heading(doc, "VIII. SYSTEM TESTING AND VALIDATION", 1)
for text in [
    (
        "Testing was conducted across functional workflows, integration points, OCR correction, and user acceptance activities. "
        "Functional tests checked whether login, catalogue search, user management, checkout, return, renewal, and OCR review "
        "behaved according to the specified use cases. Integration tests focused on API-to-database behavior and status updates, "
        "especially where a circulation action must update book-copy availability and create a loan transaction."
    ),
    (
        "The OCR test case was treated as a critical validation scenario. When the engine produced a mistaken value, the expected "
        "behavior was that the librarian could compare the result with the source image, edit the field, save the correction, and "
        "commit the verified record. This human-in-the-loop step reduces the risk of permanently storing incorrect OCR output."
    ),
]:
    p = add_para(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_table(
    doc,
    "Table IV: OCR Verification Test Case",
    ["Item", "Description"],
    [
        ("Test case", "Verify and correct OCR data before committing a ledger row."),
        ("Input", "Uploaded ledger image with an OCR typo, for example 'Sej@rah'."),
        ("Expected result", "Uploaded image and extracted fields are shown side by side; the librarian edits the field to 'Sejarah'."),
        ("Pass condition", "Corrected data is saved and committed to the catalogue only after staff review."),
    ],
    [0.9, 2.1],
)

add_heading(doc, "IX. ACHIEVEMENT OF PROJECT OBJECTIVES", 1)
p = add_para(
    doc,
    (
        "The completed project achieved its three objectives. The system analysis and design addressed manual record keeping and "
        "borrowing pain points. The implementation delivered a working web-based library system with OCR-assisted ledger "
        "digitization, barcode circulation, role-based access, and bilingual interfaces. Evaluation activities confirmed that "
        "core modules functioned as intended and that the OCR workflow was improved through staff review before final database "
        "commitment."
    ),
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_table(
    doc,
    "Table V: Objective Achievement Summary",
    ["Objective", "Achievement"],
    [
        ("Analyze and design", "Stakeholder interviews, SRS, SDD, workflow, use case, ERD, and interface designs were completed."),
        ("Develop system", "React, Flask, Supabase, OCR review, barcode circulation, roles, loans, and bilingual pages were implemented."),
        ("Evaluate system", "Functional tests, integration checks, OCR correction testing, and user acceptance feedback were used."),
    ],
    [0.95, 2.05],
)

add_heading(doc, "X. FUTURE WORK", 1)
for text in [
    (
        "Future development can extend the OCR pipeline into a cloud-hosted or dedicated-server workflow so digitization can be used "
        "without relying on one local scanning station. A reporting module could also provide insights on borrowing trends, overdue "
        "rates, popular titles, and class-level reading activity. These analytics would support collection development and school "
        "reading programs."
    ),
    (
        "Additional improvements include integration with the school's student information system, automated reminders for due dates, "
        "reservation features, richer public catalogue functions, and stronger administrative reports. These enhancements would "
        "build on the current architecture while keeping the system aligned with actual school operations."
    ),
]:
    p = add_para(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_heading(doc, "XI. CONCLUSION", 1)
p = add_para(
    doc,
    (
        "ePustaka Munshi demonstrates how a school-specific library management system can combine ordinary operational features "
        "with a focused digitization workflow for historical handwritten records. By integrating a searchable catalogue, role-based "
        "administration, barcode circulation, and OCR-assisted ledger review, the system reduces reliance on fragile manual records "
        "while preserving the evidence needed for accurate catalogue migration. The project also shows that modern web tools and "
        "vision-based OCR can be applied pragmatically in a Malaysian secondary-school library when the design keeps human review, "
        "workflow fit, and maintainability at the center."
    ),
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_heading(doc, "ACKNOWLEDGMENT", 1)
p = add_para(
    doc,
    (
        "The author expresses sincere appreciation to Dr. Siti Nur Khadijah Aishah Binti Ibrahim for supervision and guidance, and "
        "to Cikgu Farah Intan Nailah Binti Adanan and the SMK Abdullah Munshi library stakeholders for providing operational input, "
        "library context, and feedback throughout the project."
    ),
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_heading(doc, "REFERENCES", 1)
references = [
    "Alkhidmah Journal. (2023). Digital library management system study. Journal of Information Systems and Technology, 4(2), 45-58. https://ejurnalqarnain.stisnq.ac.id/index.php/ALKHIDMAH/article/view/501",
    "CDS Online. (n.d.). Open source library systems & digitisation. Retrieved October 15, 2025, from https://cdsol.com.my/web/open-source-library-systems-digitisation/",
    "Francis, S. A., & Sangeetha, M. (2025). A comparison study on optical character recognition models in mathematical equations and in any language. Results in Control and Optimization, 18, 100532. https://doi.org/10.1016/j.rico.2025.100532",
    "Humphries, M., Leddy, L. C., Downton, Q., Legace, M., McConnell, J., Murray, I., & Spence, E. (2024). Unlocking the archives: Using large language models to transcribe handwritten historical documents. arXiv. https://doi.org/10.48550/arXiv.2411.03340",
    "Imam Yuadi, Sigh, A. R., & Ullin Nihaya. (2024). Text recognition for library collection in different light conditions. TEM Journal, 266-276. https://doi.org/10.18421/tem131-28",
    "Lesmana, C., Santoso, D., & Mannanan, D. L. (2025). Development of web-based library system at SMP Negeri 3 Sungai Kakap. International Journal of Multi Discipline Science, 8(1), 43-51. https://doi.org/10.26737/ij-mds.v8i1.6719",
    "Libsys Malaysia. (n.d.). Library management system solutions. Retrieved November 2, 2025, from https://www.libsys.my/",
    "Memon, J., Sami, M., Khan, R. A., & Uddin, M. (2020). Handwritten optical character recognition (OCR): A comprehensive systematic literature review (SLR). IEEE Access, 8, 142642-142668. https://doi.org/10.1109/access.2020.3012542",
    "Multazam, A. E., Qashlim, A., & Sarjan, M. (2023). Image processing technology in book metadata extraction system using optical character recognition (OCR). Jurnal Sisfotek Global, 13(1), 1-7. https://doi.org/10.38101/sisfotek.v13i1.865",
    "Nazeem, M. (n.d.). Open-source OCR libraries: A comprehensive study for low resource language. https://aclanthology.org/2024.icon-1.48.pdf",
    "Putri Aisyah Hasibuan, Rahmat Fadhli, & Miftahunnisa' Igiriza. (2023). Redefining school libraries for the digital age: Developing comprehensive digital collection strategies. Jurnal Manajemen Pendidikan Jurnal Ilmiah Administrasi Manajemen Dan Kepemimpinan Pendidikan, 5(1), 58-68. https://doi.org/10.21831/jump.v5i1.60752",
    "Shi, Y., Peng, D., Liao, W., Lin, Z., Chen, X., Liu, C., Zhang, Y., & Jin, L. (2023). Exploring OCR capabilities of GPT-4V(ision): A quantitative and in-depth evaluation. arXiv. https://doi.org/10.48550/arXiv.2310.16809",
    "Smith, R. (2021). Tesseract.js documentation: Pure JavaScript OCR for 62 languages. Retrieved December 10, 2025, from https://tesseract.projectnaptha.com/",
    "Supabase. (2024). Supabase documentation: The open source Firebase alternative. Retrieved November 20, 2025, from https://supabase.com/docs",
    "Vamvakas, G., Gatos, B., Stamatopoulos, N., & Perantonis, S. J. (2008). A complete optical character recognition methodology for historical documents. Proceedings of the Eighth IAPR Workshop on Document Analysis Systems (DAS 2008), 525-532. IEEE. https://doi.org/10.1109/DAS.2008.73",
    "Wang, J. (2023). A study of the OCR development history and directions of development. Highlights in Science Engineering and Technology, 72, 409-415. https://doi.org/10.54097/bm665j77",
]
for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"[{i}] {ref}")
    set_run_font(r, size=8.5)

OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "ePustaka Munshi: Smart Library System with OCR Ledger Digitization"
doc.core_properties.author = "Wan Zafirzan Bin Wan Tarmizan"
doc.save(OUT_PATH)
print(OUT_PATH)
