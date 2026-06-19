from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT_PATH = Path(
    r"C:\Users\wanza\OneDrive - Universiti Teknologi Malaysia (UTM)\UTM\Final Year Project\FINAL"
    r"\ePustaka Munshi Short Paper - Two Column Draft.docx"
)

FONT = "Times New Roman"


def set_run_font(run, size=10, bold=None, italic=None, small_caps=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if small_caps is not None:
        run.font.small_caps = small_caps


def set_columns(section, num=1, space=360):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space))


def configure_section(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0


def add_para(doc, text="", size=10, bold=False, italic=False, align=None, after=4, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def add_section_heading(doc, roman, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    heading_text = f"{roman}.  {title.upper()}" if roman else title.upper()
    r = p.add_run(heading_text)
    set_run_font(r, size=10, small_caps=True)
    return p


def add_placeholder(doc, label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"\n\n[IMAGE HERE - {label}]\n\n")
    set_run_font(r, size=8.8, bold=True)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=8.5)
    return p


def add_table_placeholder(doc, caption, rows):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(caption)
    set_run_font(r, size=8.6, bold=True)

    block = doc.add_paragraph()
    block.alignment = WD_ALIGN_PARAGRAPH.LEFT
    block.paragraph_format.left_indent = Inches(0.08)
    block.paragraph_format.right_indent = Inches(0.08)
    block.paragraph_format.space_before = Pt(0)
    block.paragraph_format.space_after = Pt(4)
    for idx, row in enumerate(rows):
        run = block.add_run(row)
        set_run_font(run, size=7.7, bold=idx == 0)
        if idx != len(rows) - 1:
            block.add_run("\n")
    return block


def body_para(doc, text):
    p = add_para(doc, text, size=10, after=5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.18)
    return p


def ref_para(doc, text):
    p = add_para(doc, text, size=7.8, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def author_pair_line(doc, left, right, size=9.3, bold=False, after=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.tab_stops.add_tab_stop(Inches(1.72), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(5.18), WD_TAB_ALIGNMENT.CENTER)
    p.add_run().add_tab()
    left_run = p.add_run(left)
    set_run_font(left_run, size=size, bold=bold)
    p.add_run().add_tab()
    right_run = p.add_run(right)
    set_run_font(right_run, size=size, bold=bold)
    return p


doc = Document()
configure_styles(doc)
for section in doc.sections:
    configure_section(section)
    set_columns(section, 1)

add_para(
    doc,
    "ePustaka Munshi: Smart Library System with OCR Ledger Digitization",
    size=21,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    after=14,
)

author_pair_line(
    doc,
    "Wan Zafirzan Bin Wan Tarmizan",
    "Dr. Siti Nur Khadijah Aishah Binti Ibrahim",
    size=10.2,
    bold=True,
)
author_pair_line(
    doc,
    "Malaysia-Japan International Institute of Technology",
    "Malaysia-Japan International Institute of Technology",
    size=8.8,
)
author_pair_line(doc, "Universiti Teknologi Malaysia", "Universiti Teknologi Malaysia", size=9.1)
author_pair_line(doc, "Email: [your email here]", "Email: [supervisor email here]", size=9.1, after=10)

abstract = (
    "Abstract--School libraries that still depend on handwritten accession ledgers and manual circulation books face persistent "
    "problems in catalogue search, book-status tracking, borrowing control, and long-term preservation of historical records. "
    "This paper presents ePustaka Munshi, a smart library system developed for SMK Abdullah Munshi to modernize a paper-based "
    "secondary-school library workflow. The system integrates role-based access, book and copy cataloguing, student and member "
    "management, barcode-based borrowing and returning, bilingual interfaces, and an OCR-assisted workflow for digitizing the "
    "Buku Induk ledger. The implementation uses a React and TypeScript frontend, a Flask REST API, and Supabase PostgreSQL as "
    "the managed relational data layer. Handwritten ledger extraction uses Claude Vision as the primary vision-language model, "
    "with Tesseract retained as a traditional OCR baseline and page-orientation aid. Developed through Agile Scrum, the system "
    "provides a searchable catalogue, faster circulation, and a human-verified digitization process that protects legacy library "
    "records from direct blind import."
)
abstract_p = add_para(doc, abstract, size=9.4, bold=True, after=6)
abstract_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

keywords = (
    "Keywords--smart library system; OCR ledger digitization; handwritten document transcription; barcode circulation; "
    "React; Flask; Supabase; school library."
)
kw_p = add_para(doc, keywords, size=9.2, italic=True, after=10)
kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

body_section = doc.add_section(WD_SECTION.CONTINUOUS)
configure_section(body_section)
set_columns(body_section, 2, space=330)

add_section_heading(doc, "I", "Introduction")
body_para(
    doc,
    "Library resource management is a daily operational requirement for supporting reading culture, reference use, and school administration. "
    "At SMK Abdullah Munshi, the library manages more than 7,000 books while still relying on physical ledgers, manual borrowing records, "
    "and fragmented spreadsheet data. Important information such as titles, authors, publishers, Dewey Decimal Classification numbers, "
    "accession records, and copy availability is therefore difficult to search and maintain consistently."
)
body_para(
    doc,
    "The manual process creates three connected problems. First, the handwritten Buku Induk is valuable but fragile, and searching it requires "
    "physical inspection. Second, borrowing and returning require repeated writing and manual status updates at the circulation counter. Third, "
    "previous digitization efforts did not provide validation, duplication checking, role-based permissions, or a reliable path for migrating old "
    "ledger rows into a live catalogue. Similar studies have shown that school-library modernization depends on centralized metadata, searchable "
    "records, and practical digital collection strategies [1], [6], [11]."
)
body_para(
    doc,
    "This project therefore developed ePustaka Munshi, a web-based library management system with OCR ledger digitization. The goal is not only "
    "to replace paper with a database, but to fit the actual operating environment of a Malaysian secondary-school library, where teachers supervise, "
    "library prefects support counter work, students need simple catalogue access, and old handwritten records must be preserved before they degrade."
)

add_section_heading(doc, "II", "Background and Related Work")
body_para(
    doc,
    "Candidate solutions were reviewed against the needs of SMK Abdullah Munshi. Mature integrated library systems such as KOHA provide broad library "
    "functionality, while commercial systems offer managed support and specialized modules [2], [7]. However, these platforms can be too heavy for a "
    "small school library that needs fast deployment, low maintenance, and workflow-specific support. Spreadsheet-based approaches are easy to begin, "
    "but they do not enforce circulation rules, audit trails, or reliable catalogue migration."
)
body_para(
    doc,
    "The digitization component required special attention because handwritten documents remain difficult for classical OCR. Traditional OCR and image "
    "processing remain useful for printed or semi-structured text [5], [9], [13], [15], [16], but handwriting introduces variation in ink, spacing, "
    "orientation, language, and writing style. Recent research on vision-capable large language models indicates that such models can transcribe and "
    "structure historical handwritten records more effectively than conventional OCR in selected archival settings [4], and GPT-4V has also been "
    "quantitatively evaluated for OCR-like tasks [12]. These findings support a hybrid approach that combines model extraction with staff verification."
)
add_table_placeholder(
    doc,
    "Table I. Comparison of Existing Library Digitization Alternatives",
    [
        "Alternative | Implication for SMK Abdullah Munshi",
        "KOHA or full ILS | Powerful, but requires more administration than the school context can easily sustain.",
        "Legacy desktop system | Familiar school-library category, but device-bound and risky for long-term maintenance.",
        "Spreadsheet inventory | Simple, but weak for circulation status, permissions, validation, and auditability.",
        "ePustaka Munshi | Tailored web workflow with OCR review, barcode circulation, roles, and managed cloud data."
    ],
)

add_section_heading(doc, "III", "Methodology")
body_para(
    doc,
    "Agile Scrum was selected because the project involved evolving stakeholder expectations and uncertain OCR performance. The work was organized "
    "into four functional sprints: user and role management, book catalogue and inventory, ledger digitization, and borrowing or returning management. "
    "Each sprint produced a functional increment that could be checked against the real library workflow."
)
body_para(
    doc,
    "Requirements were gathered through stakeholder discussion with the school librarian and review of existing library artifacts, including the Buku "
    "Induk, borrowing records, and spreadsheet inventory. The findings were converted into functional requirements, non-functional requirements, use "
    "case descriptions, workflow diagrams, database entities, interface designs, and test cases. This process kept the system grounded in school practice "
    "rather than in generic library software assumptions."
)
add_placeholder(doc, "Figure 1. Agile Scrum Development Lifecycle")
add_caption(doc, "Fig. 1. Agile Scrum development lifecycle used to structure iterative system delivery.")

add_section_heading(doc, "IV", "System Analysis")
body_para(
    doc,
    "The main actors are Administrator, Librarian, Library Prefect, and Student. Administrators manage global settings, roles, and higher-risk operations. "
    "Librarians manage book records, student data, OCR jobs, and circulation monitoring. Library Prefects assist at the counter by scanning books for "
    "borrowing and returning. Students use the portal to view catalogue records, active loans, overdue status, and borrowing history."
)
body_para(
    doc,
    "The key requirement is the OCR verification workflow. Since handwritten extraction is imperfect, the system deliberately avoids direct automatic "
    "import into the catalogue. Instead, staff upload a ledger image or PDF, review extracted rows, correct field values against the source image, and "
    "commit only verified records. This human-in-the-loop pattern reduces the risk of permanently storing OCR mistakes."
)
add_table_placeholder(
    doc,
    "Table II. Core Requirement Summary",
    [
        "Category | Requirement",
        "Authentication | Route users to the correct role-based interface after secure login.",
        "Inventory | Add, edit, search, filter, and manage book and copy records.",
        "Digitization | Upload ledger pages, extract rows, correct OCR values, and commit verified data.",
        "Circulation | Scan barcodes for borrowing, returning, renewal, overdue review, and status updates.",
        "Reliability | Prevent low-confidence OCR output from entering the permanent catalogue without review."
    ],
)
add_placeholder(doc, "Figure 2. Use Case Diagram for ePustaka Munshi")
add_caption(doc, "Fig. 2. Use case design showing role-specific access for library stakeholders.")

add_section_heading(doc, "V", "System Design and Implementation")
body_para(
    doc,
    "The system follows a three-tier web architecture. The presentation layer is a React single-page application with module-based pages for "
    "authentication, catalogue browsing, user management, circulation, OCR review, and student access. The application layer is a Flask REST API that "
    "handles validation, authentication-related logic, business rules, and role permissions. The data layer is Supabase PostgreSQL, selected for managed "
    "relational persistence and cloud availability [14]."
)
body_para(
    doc,
    "The database separates access control, membership, bibliographic data, physical copies, loans, OCR evidence, and digitized ledger archives. Book "
    "records store bibliographic metadata, while book-copy records store accession numbers, barcodes, condition, shelf location, and availability status. "
    "Loan records connect members, copies, handlers, due dates, renewals, and transaction status. OCR jobs preserve uploaded evidence, extracted rows, "
    "confidence values, correction history, and final commit status."
)
body_para(
    doc,
    "The interface prioritizes fast operation at the library counter. The barcode field remains easy to focus so a USB scanner can behave like keyboard "
    "input, while immediate visual feedback confirms a successful checkout or return. The OCR screen works as a verification workspace, showing the source "
    "image and extracted rows so the librarian can correct title, author, publisher, accession number, and classification values before import."
)
add_placeholder(doc, "Figure 3. High-Level System Workflow Swimlane Diagram")
add_caption(doc, "Fig. 3. High-level workflow showing how catalogue, OCR, and circulation activities connect.")
add_placeholder(doc, "Figure 4. Entity Relationship Diagram for ePustaka Munshi")
add_caption(doc, "Fig. 4. Entity-relationship design showing the main persistence structure.")
add_placeholder(doc, "Figure 5. OCR Digitization Review Interface")
add_caption(doc, "Fig. 5. OCR review interface where staff verify extracted ledger rows before committing records.")
add_placeholder(doc, "Figure 6. Barcode Checkout and Return Interface")
add_caption(doc, "Fig. 6. Barcode circulation interface for fast borrowing and returning at the counter.")

add_section_heading(doc, "VI", "Testing and Evaluation")
body_para(
    doc,
    "Testing covered functional workflows, integration behavior, OCR correction, and user acceptance activities. Functional tests checked login, catalogue "
    "search, user management, checkout, return, renewal, and OCR review. Integration checks focused on API-to-database behavior, especially where a "
    "circulation action must update both the loan transaction and the copy availability status."
)
body_para(
    doc,
    "The OCR scenario was treated as a critical validation case. When the engine produced an incorrect value, such as a misspelled title or uncertain "
    "classification field, the expected behavior was that the librarian could compare the result with the source image, edit the field, save the correction, "
    "and commit the verified record. This confirmed that OCR is used as an assistant rather than as an unchecked authority."
)
add_table_placeholder(
    doc,
    "Table III. OCR Verification Test Case",
    [
        "Item | Description",
        "Input | Ledger image with an OCR typo or missing bibliographic field.",
        "Expected result | Source image and extracted fields are shown side by side for review.",
        "Correction | Librarian edits the field and saves the reviewed row.",
        "Pass condition | Corrected data is committed only after staff verification."
    ],
)

add_section_heading(doc, "VII", "Results and Discussion")
body_para(
    doc,
    "The completed system achieved the project objectives. The analysis and design activities documented the manual workflow and translated the library "
    "pain points into software requirements. The implementation delivered a working web system with role-based access, catalogue management, copy tracking, "
    "barcode circulation, OCR-assisted ledger digitization, and bilingual interface support. The evaluation confirmed that the core workflows behaved as "
    "intended and that OCR output could be corrected before permanent storage."
)
body_para(
    doc,
    "The most important design decision is the separation between extraction and commitment. This protects catalogue quality while still reducing the "
    "manual burden of retyping old ledger pages. The architecture also avoids the weakness of a single-computer desktop deployment because the database is "
    "hosted, relational, and easier to back up. At the same time, the system remains practical for a school because counter operations are simple and "
    "library prefects can assist under limited permissions."
)

add_section_heading(doc, "VIII", "Conclusion and Future Work")
body_para(
    doc,
    "ePustaka Munshi demonstrates how a school-specific library management system can combine ordinary operational features with a focused digitization "
    "workflow for historical handwritten records. By integrating a searchable catalogue, role-based administration, barcode circulation, and OCR-assisted "
    "ledger review, the system reduces reliance on fragile manual records while preserving the evidence needed for accurate catalogue migration."
)
body_para(
    doc,
    "Future work can move more of the OCR pipeline into a cloud-hosted or dedicated-server workflow, add borrowing analytics, support automated due-date "
    "reminders, provide reservation features, and integrate with school student information systems. These enhancements can build on the current architecture "
    "while keeping the system aligned with actual library operations."
)

add_section_heading(doc, "", "Acknowledgment")
body_para(
    doc,
    "The author expresses sincere appreciation to Dr. Siti Nur Khadijah Aishah Binti Ibrahim for supervision and guidance, and to Cikgu Farah Intan Nailah "
    "Binti Adanan and the SMK Abdullah Munshi library stakeholders for operational input, library context, and project feedback."
)

add_section_heading(doc, "", "References")
references = [
    "Alkhidmah Journal. (2023). Digital library management system study. Journal of Information Systems and Technology, 4(2), 45-58. https://ejurnalqarnain.stisnq.ac.id/index.php/ALKHIDMAH/article/view/501",
    "CDS Online. (n.d.). Open source library systems & digitisation. Retrieved October 15, 2025, from https://cdsol.com.my/web/open-source-library-systems-digitisation/",
    "Francis, S. A., & Sangeetha, M. (2025). A comparison study on optical character recognition models in mathematical equations and in any language. Results in Control and Optimization, 18, 100532. https://doi.org/10.1016/j.rico.2025.100532",
    "Humphries, M., Leddy, L. C., Downton, Q., Legace, M., McConnell, J., Murray, I., & Spence, E. (2024). Unlocking the archives: Using large language models to transcribe handwritten historical documents. arXiv. https://doi.org/10.48550/arXiv.2411.03340",
    "Imam Yuadi, Sigh, A. R., & Ullin Nihaya. (2024). Text recognition for library collection in different light conditions. TEM Journal, 266-276. https://doi.org/10.18421/tem131-28",
    "Lesmana, C., Santoso, D., & Mannanan, D. L. (2025). Development of web-based library system at SMP Negeri 3 Sungai Kakap. International Journal of Multi Discipline Science, 8(1), 43-51. https://doi.org/10.26737/ij-mds.v8i1.6719",
    "Libsys Malaysia. (n.d.). Library management system solutions. Retrieved November 2, 2025, from https://www.libsys.my/",
    "Memon, J., Sami, M., Khan, R. A., & Uddin, M. (2020). Handwritten optical character recognition (OCR): A comprehensive systematic literature review (SLR). IEEE Access, 8, 142642-142668. https://doi.org/10.1109/access.2020.3012542",
    "Multazam, A. E., Qashlim, A., & Sarjan, M. (2023). Image processing technology in book metadata extraction system using optical character recognition (OCR). Jurnal Sisfotek Global, 13(1), 1-7. https://doi.org/10.38101/sisfotek.v13i1.865",
    "Nazeem, M. (n.d.). Open-source OCR libraries: A comprehensive study for low resource language. https://aclanthology.org/2024.icon-1.48.pdf",
    "Putri Aisyah Hasibuan, Rahmat Fadhli, & Miftahunnisa' Igiriza. (2023). Redefining school libraries for the digital age: Developing comprehensive digital collection strategies. Jurnal Manajemen Pendidikan Jurnal Ilmiah Administrasi Manajemen Dan Kepemimpinan Pendidikan, 5(1), 58-68. https://doi.org/10.21831/jump.v5i1.60752",
    "Shi, Y., Peng, D., Liao, W., Lin, Z., Chen, X., Liu, C., Zhang, Y., & Jin, L. (2023). Exploring OCR capabilities of GPT-4V(ision): A quantitative and in-depth evaluation. arXiv. https://doi.org/10.48550/arXiv.2310.16809",
    "Smith, R. (2021). Tesseract.js documentation: Pure JavaScript OCR for 62 languages. Retrieved December 10, 2025, from https://tesseract.projectnaptha.com/",
    "Supabase. (2024). Supabase documentation: The open source Firebase alternative. Retrieved November 20, 2025, from https://supabase.com/docs",
    "Vamvakas, G., Gatos, B., Stamatopoulos, N., & Perantonis, S. J. (2008). A complete optical character recognition methodology for historical documents. Proceedings of the Eighth IAPR Workshop on Document Analysis Systems (DAS 2008), 525-532. IEEE. https://doi.org/10.1109/DAS.2008.73",
    "Wang, J. (2023). A study of the OCR development history and directions of development. Highlights in Science Engineering and Technology, 72, 409-415. https://doi.org/10.54097/bm665j77",
]
for index, reference in enumerate(references, 1):
    ref_para(doc, f"[{index}] {reference}")

OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "ePustaka Munshi: Smart Library System with OCR Ledger Digitization"
doc.core_properties.author = "Wan Zafirzan Bin Wan Tarmizan"
doc.save(OUT_PATH)
print(OUT_PATH)
